import re

from docx import Document

from modules.commission.converters.base import (
    DocumentConverter,
    ConversionResult,
    EMPTY_TEXT_THRESHOLD,
    THIN_TEXT_WARNING_THRESHOLD,
    build_conversion_metadata,
)


class DocxConverter(DocumentConverter):
    """Convert Word (.docx) documents to plain text using python-docx."""

    def convert(self, file_path: str, **kwargs) -> ConversionResult:
        """Extract text from a .docx file.

        Extracts paragraph text and table content in document order.

        Args:
            file_path: Path to the .docx file.

        Returns:
            ConversionResult with extracted text.
        """
        doc = Document(file_path)
        body_parts = []

        # Iterate through paragraphs and tables in document order
        for block in self._iter_block_items(doc):
            block_text = self._extract_block_text(block)
            if block_text:
                body_parts.append(block_text)

        header_footer_parts = self._extract_header_footer_parts(doc)
        parts = body_parts + header_footer_parts

        full_text = "\n\n".join(parts)
        body_text = "\n\n".join(body_parts).strip()
        header_footer_text = "\n\n".join(header_footer_parts).strip()

        warnings = []
        total_stripped_length = len(full_text.strip())
        if total_stripped_length >= EMPTY_TEXT_THRESHOLD and len(body_text) < THIN_TEXT_WARNING_THRESHOLD:
            warnings.append((
                "docx_thin_body",
                "DOCX body text is unusually thin compared with the full extracted document.",
            ))

        metadata = build_conversion_metadata(
            full_text,
            stats={
                "body_chars": len(body_text),
                "header_footer_chars": len(header_footer_text),
                "header_footer_blocks": len(header_footer_parts),
            },
            warnings=warnings,
        )

        return ConversionResult(
            text=full_text,
            page_count=None,
            pages_after_dedup=None,
            metadata=metadata,
        )

    def _iter_block_items(self, doc):
        """Yield paragraphs and tables in document order."""
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn

        for element in doc.element.body:
            if element.tag == qn("w:p"):
                yield Paragraph(element, doc)
            elif element.tag == qn("w:tbl"):
                yield Table(element, doc)

    def _extract_block_text(self, block):
        if hasattr(block, "rows"):
            return self._extract_table(block)
        if hasattr(block, "text"):
            return self._extract_paragraph(block)
        return ""

    def _extract_paragraph(self, paragraph):
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = (getattr(paragraph.style, "name", "") or "").lower()
        if "list bullet" in style_name or style_name == "list paragraph":
            return text if re.match(r"^[-*•]\s", text) else f"- {text}"
        if "list number" in style_name:
            return text if re.match(r"^\d+[.)]\s", text) else f"1. {text}"
        return text

    def _extract_table(self, table):
        """Extract text from a python-docx Table object."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        return "\n".join(rows) if rows else ""

    def _extract_header_footer_parts(self, doc):
        seen_parts = set()
        parts = []

        for section in doc.sections:
            for story_part in (section.header, section.footer):
                for paragraph in story_part.paragraphs:
                    text = self._extract_paragraph(paragraph)
                    if text and text not in seen_parts:
                        seen_parts.add(text)
                        parts.append(text)
                for table in story_part.tables:
                    table_text = self._extract_table(table)
                    if table_text and table_text not in seen_parts:
                        seen_parts.add(table_text)
                        parts.append(table_text)

        return parts
